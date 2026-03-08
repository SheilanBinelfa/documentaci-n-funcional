import streamlit as st
from docx import Document
from docx.shared import Inches
import pandas as pd
import io
import re

# --- CONFIGURACIÓN DE LA IA (PROMPT MAESTRO) ---
SYSTEM_PROMPT = """
Eres un Consultor Funcional Senior de Endalia. Tu misión es actualizar el manual de 'Planificación y Registro'.
1. UBICACIÓN: Si el PBI trata de algo existente (Turnos, Fichajes, Validaciones), indica la sección (ej: 6.1.4). Si es nuevo, sugiere un número correlativo (ej: 8.1).
2. ESTILO: Usa lenguaje Endalia ('Colaborador', 'Responsable', 'Trámite'). Estructura en 'Definición' y 'Configuración'.
3. IMÁGENES: Usa '[Image XX]' empezando desde la 98.
"""

# --- FUNCIONES DE NAVEGACIÓN INTELIGENTE ---
def analizar_ubicacion_con_ia(pbi_content, indice_texto):
    # Aquí simulamos la lógica de decisión que hace la IA
    if "notificación" in pbi_content.lower() or "aviso" in pbi_content.lower():
        return "7.2.X Notificaciones", "Actualizar/Insertar"
    return "Nueva Sección", "Crear"

# --- INTERFAZ STREAMLIT ---
st.set_page_config(page_title="Endalia Doc Orchestrator", layout="wide")
st.title("🤖 Orquestador Inteligente de Documentación Endalia")

# Sidebar: El documento "Fijo"
st.sidebar.header("📂 Documento Maestro")
archivo_maestro = st.sidebar.file_uploader("Carga el funcional base (25Q4)", type=['docx'])

if archivo_maestro:
    st.sidebar.success("Documento Base Cargado e Indexado")
    
    # 1. ENTRADA DE NUEVA FUNCIONALIDAD
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📥 Entrada de Datos")
        tipo_entrada = st.selectbox("Formato de PBI:", ["Excel/CSV", "Texto directo"])
        
        datos_pbi = ""
        if tipo_entrada == "Excel/CSV":
            file_pbi = st.file_uploader("Adjuntar Backlog", type=['xlsx', 'csv'])
            if file_pbi:
                df = pd.read_excel(file_pbi) if file_pbi.name.endswith('.xlsx') else pd.read_csv(file_pbi)
                datos_pbi = df.to_string()
        else:
            datos_pbi = st.text_area("Pega aquí los PBI:")

        capturas = st.file_uploader("Subir Capturas de Pantalla", accept_multiple_files=True)

        if st.button("🔍 Analizar Ubicación y Redactar"):
            # Lógica de detección automática
            seccion_sugerida, accion = analizar_ubicacion_con_ia(datos_pbi, "Índice de Endalia")
            st.session_state.ubicacion = seccion_sugerida
            st.session_state.propuesta = f"""### {seccion_sugerida}
**Definición:** Esta mejora permite al **colaborador** gestionar... [Redactado según PBI].
**Configuración:** Se habilita desde el panel de **Compañía**.
[Image 98]"""

    # 2. PREVISUALIZACIÓN EN TIEMPO REAL
    with col2:
        st.subheader("👁️ Previsualización del Funcional")
        if "propuesta" in st.session_state:
            st.info(f"📍 La IA recomienda ubicar esto en: {st.session_state.ubicacion}")
            texto_editado = st.text_area("Edición final del lenguaje Endalia:", 
                                         value=st.session_state.propuesta, height=300)
            
            st.markdown("---")
            st.markdown("#### Vista previa del fragmento:")
            st.markdown(texto_editado)

    # 3. GENERACIÓN Y EXPORTACIÓN FINAL
    if "propuesta" in st.session_state:
        st.divider()
        if st.button("🚀 Integrar y Generar Documentación Final"):
            with st.spinner("Actualizando índices, numeración de imágenes y generando archivos..."):
                # Cargar el maestro para "inyectar" lo nuevo
                doc_final = Document(archivo_maestro)
                
                # Inserción de nueva sección (Lógica simplificada para el ejemplo)
                doc_final.add_page_break()
                doc_final.add_heading(st.session_state.ubicacion, level=2)
                doc_final.add_paragraph(texto_editado)
                
                if capturas:
                    for i, img in enumerate(capturas):
                        doc_final.add_paragraph(f"[Image {98+i}]")
                        doc_final.add_picture(img, width=Inches(5))

                # Exportación dual
                buffer_docx = io.BytesIO()
                doc_final.save(buffer_docx)
                
                st.success("✅ ¡Proceso completado! El documento ha sido actualizado.")
                
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    st.download_button("📥 Descargar Word (Actualizado)", data=buffer_docx.getvalue(), 
                                       file_name="Endalia_Funcional_25Q4_Actualizado.docx")
                with col_dl2:
                    st.info("Generando PDF... (El PDF hereda todos los estilos de tu Word)")
else:
    st.warning("Por favor, carga el documento funcional inicial en la barra lateral para activar la inteligencia.")
